from docx import Document

# Загружаем документ
doc = Document("all.docx")

# Вводим имя пользователя
name = input("Введите ФИО: ")

# Итерируем по всем параграфам
for paragraph in doc.paragraphs:
    # Заменяем {name} в тексте параграфа
    try:
        paragraph.text = paragraph.text.format(ФИО=name)
    except KeyError:
        continue

# Итерируем по всем таблицам
for table in doc.tables:
    # Итерируем по всем ячейкам в таблице
    for row in table.rows:
        for cell in row.cells:
            try:
                # Заменяем {name} в тексте ячейки
                cell.text = cell.text.format(ФИО=name)
            except KeyError:
                continue

# Сохраняем изменения
doc.save("example2.docx")
