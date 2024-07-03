from sys import platform
from time import strftime
from random import randint
from docx import Document as _Document
from api.models import Company

class DocError(Exception): pass

class Document:

	def __init__(self, session:str=''):
		# if len(data) != 16:
		# 	raise DocError('Not all data has been transfered!')
		if session == '':
			raise DocError('The user\'s session data has not been transmitted!')
		# self._data = data
		self._session = session

	def _collect_data(self) -> dict:
		# Мне менее больно от этого
		data = Company.objects.filter(session=self._session)
		return {
			"ФИО": data[0].fcs,
			"фио": data[0].fcs,
			"Дата протокола": strftime("%d.%m.%y"),
			"Член1": data[0].fcs_commissions1,
			"Член2": data[0].fcs_commissions2,
			"Член3": data[0].fcs_commissions3,
			"Д1": data[0].commisions_status1,
			"Д2": data[0].commisions_status2,
			"Д3": data[0].commisions_status3,
			"Должность": data[0].status,
			"Причина": data[0].reason_for_check,
			"Компания": data[0].company_name,
			"номер программы": "20", # For example
			"для кого": data[0].status+'а',
			"№": '1', # ???
			"№2": '1', # ???????
			"группа": data[0].electrical_safety_group,
			"группаП": data[0].electrical_safety_group,
			"Дата ЭБ": f"{randint(1, 30)}.{randint(1, 12)}{randint(2025, 2027)}", # For example
			"ЭНФ": "(энф?)",
			"Инструктаж": data[0].reason_for_check,
			"Дата ПЭБ": f"{randint(1, 30)}.{randint(1, 12)}{randint(2025, 2027)}", # For example
			"Стаж": data[0].work_exp
		}

	def _replacement(self, words:dict, filename:str) -> str:
		doc = _Document(filename)
		for paragraph in doc.paragraphs:
			try:
				paragraph.text = paragraph.text.format(**words)
			except KeyError:
				continue
		for table in doc.tables:
			for row in table.rows:
				for cell in row.cells:
					try:
						cell.text = cell.text.format(**words)
					except KeyError:
						continue
		doc.save(filename.split('.')[0]+self._session+'.docx')
		return filename.split('.')[0]+self._session+'.docx'

	def get(self, filename:str) -> str:
		words = self._collect_data()
		return self.convert_to_pdf(self._replacement(words, filename))

	def convert_to_pdf(self, filename:str) -> str:
		pdf_filename = filename.split('.')[0]+'.pdf'
		if platform == 'win32':
			from docx2pdf import convert
			convert(filename, pdf_filename)
		else: # For Linux
			from os import rename, remove
			from subprocess import run

			run(['librepffoce', '--headless', '--convert-to', 'pdf', filename, '--outdir', 'tmp'], check=True)
			rename(f'tmp/{pdf_filename}', pdf_filename)
			remove(f'tmp/pdf_filename')
		return pdf_filename